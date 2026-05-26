"""
IT감사 보고서 자동 생성기
- Claude API로 위반 데이터 기반 자연어 분석 문장 생성
- Excel: 규칙별 위반 현황표 + 도메인별 요약
- Word:  표지 + 총평 + 도메인별 상세 + 시정조치

입력:  data/processed/violations_summary.csv
출력:  data/processed/report/IT감사보고서_YYYYMMDD.xlsx
       data/processed/report/IT감사보고서_YYYYMMDD.docx
"""

import os
import json
import pandas as pd
from datetime import datetime
from dotenv import dotenv_values

# ── 경로 설정 ────────────────────────────────────────────────
BASE_DIR   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SUMMARY_PATH = f"{BASE_DIR}/data/processed/violations_summary.csv"
OUT_DIR    = f"{BASE_DIR}/data/processed/report"

# ── Claude API 설정 ──────────────────────────────────────────
_env = dotenv_values(os.path.join(BASE_DIR, ".env"))
_api_key = _env.get("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_API_KEY")


# ══════════════════════════════════════════════════════════════
# 1. Claude API - 자연어 분석 생성
# ══════════════════════════════════════════════════════════════
def generate_analysis(df: pd.DataFrame) -> dict:
    """
    violations_summary를 Claude에게 전달해서
    - 총평 (executive_summary)
    - 도메인별 분석 (domain_analysis: 접근통제/변경관리/운영통제)
    - 핵심 시정조치 (key_actions)
    를 자연어로 생성
    """
    from anthropic import Anthropic
    client = Anthropic(api_key=_api_key)

    # 위반 데이터 요약 (API 전달용)
    total     = len(df)
    violated  = (df["yn_violation"] == "Y").sum()
    high_viol = df[(df["yn_violation"] == "Y") & (df["severity"] == "HIGH")]

    domain_summary = df.groupby("audit_domain").agg(
        전체=("rule_id", "count"),
        위반=("yn_violation", lambda x: (x == "Y").sum()),
        총건수=("violation_count", "sum")
    ).reset_index()

    top_violations = df[df["yn_violation"] == "Y"].nlargest(10, "violation_count")[
        ["rule_nm", "audit_domain", "severity", "violation_count", "source_law", "remediation"]
    ]

    data_summary = f"""
[점검 개요]
- 점검 기준일: {datetime.now().strftime('%Y년 %m월 %d일')}
- 총 점검 규칙: {total}개
- 위반 탐지: {violated}개 ({violated/total*100:.1f}%)
- 이상 없음: {total - violated}개

[도메인별 위반 현황]
{domain_summary.to_string(index=False)}

[HIGH 위반 규칙 ({len(high_viol)}개)]
{high_viol[['rule_nm','audit_domain','violation_count','source_law']].to_string(index=False)}

[위반 건수 TOP 10]
{top_violations.to_string(index=False)}
"""

    prompt = f"""당신은 금융권 IT감사 전문가입니다.
아래 Rule 엔진 점검 결과를 바탕으로 IT감사 보고서에 들어갈 분석 내용을 작성해주세요.

{data_summary}

다음 항목을 JSON 형식으로 출력하세요:
{{
  "executive_summary": "경영진 보고용 총평 (3~5문장, 핵심 위반 현황·위험도·조치 방향 포함)",
  "domain_analysis": {{
    "접근통제": "접근통제 분야 분석 (3~4문장, 주요 위반 유형·위험 원인·영향 포함)",
    "변경관리": "변경관리 분야 분석 (3~4문장)",
    "운영통제": "운영통제 분야 분석 (3~4문장)"
  }},
  "key_actions": [
    "즉시 조치 (1순위): ...",
    "단기 조치 (1개월 내): ...",
    "중기 조치 (3개월 내): ...",
    "장기 조치 (6개월 내): ..."
  ],
  "risk_opinion": "전반적 IT통제 위험 수준에 대한 종합 의견 (2~3문장)"
}}

반드시 JSON만 출력하세요."""

    print("  Claude API 분석 생성 중...")
    response = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    raw = response.content[0].text.strip()
    if "```json" in raw:
        raw = raw.split("```json")[1].split("```")[0].strip()
    elif "```" in raw:
        raw = raw.split("```")[1].split("```")[0].strip()

    return json.loads(raw)


# ══════════════════════════════════════════════════════════════
# 2. Excel 보고서 생성
# ══════════════════════════════════════════════════════════════
def generate_excel(df: pd.DataFrame, analysis: dict, out_path: str):
    """
    시트 구성:
    - 요약: 도메인별 위반 현황 + 총평
    - 전체규칙: 70개 규칙 전체 목록
    - 위반상세: 위반 탐지된 규칙만
    - HIGH위반: HIGH 심각도 위반만
    """
    from openpyxl import Workbook
    from openpyxl.styles import (Font, PatternFill, Alignment,
                                  Border, Side, GradientFill)
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, Reference

    wb = Workbook()

    # ── 색상 정의 ──────────────────────────────────────────────
    COLOR = {
        "header_blue":  "1F4E79",
        "header_light": "BDD7EE",
        "red":          "FF0000",
        "orange":       "FF8C00",
        "green":        "00B050",
        "yellow":       "FFFF00",
        "gray_bg":      "F2F2F2",
        "white":        "FFFFFF",
        "dark_text":    "000000",
    }

    sev_fill = {
        "HIGH":   PatternFill("solid", fgColor="FF0000"),
        "MEDIUM": PatternFill("solid", fgColor="FFA500"),
        "LOW":    PatternFill("solid", fgColor="00B050"),
    }
    viol_fill = {
        "Y": PatternFill("solid", fgColor="FFE0E0"),
        "N": PatternFill("solid", fgColor="E0FFE0"),
    }

    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def header_style(cell, bg=COLOR["header_blue"], fg=COLOR["white"], size=11, bold=True):
        cell.font = Font(bold=bold, color=fg, size=size)
        cell.fill = PatternFill("solid", fgColor=bg)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    def data_style(cell, bold=False, align="left"):
        cell.font = Font(bold=bold, size=10)
        cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
        cell.border = border

    def set_col_width(ws, col, width):
        ws.column_dimensions[get_column_letter(col)].width = width

    # ────────────────────────────────────────────────────────────
    # 시트 1: 요약
    # ────────────────────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "요약"
    ws1.row_dimensions[1].height = 40

    # 제목
    ws1.merge_cells("A1:H1")
    title_cell = ws1["A1"]
    title_cell.value = "IT감사 Rule 엔진 점검 결과 보고서"
    title_cell.font = Font(bold=True, size=16, color=COLOR["white"])
    title_cell.fill = PatternFill("solid", fgColor=COLOR["header_blue"])
    title_cell.alignment = Alignment(horizontal="center", vertical="center")

    # 점검 기준일
    ws1["A2"] = f"점검 기준일: {datetime.now().strftime('%Y년 %m월 %d일')}"
    ws1["A2"].font = Font(size=10, color="595959")
    ws1.merge_cells("A2:H2")

    # KPI 박스
    total    = len(df)
    violated = (df["yn_violation"] == "Y").sum()
    high_cnt = ((df["yn_violation"] == "Y") & (df["severity"] == "HIGH")).sum()

    kpi_data = [
        ("총 점검 규칙", f"{total}개", COLOR["header_blue"]),
        ("위반 탐지",    f"{violated}개 ({violated/total*100:.1f}%)", "C00000"),
        ("HIGH 위반",    f"{high_cnt}개", "FF4500"),
        ("이상 없음",    f"{total-violated}개", "00B050"),
    ]
    kpi_cols = [1, 3, 5, 7]
    ws1.row_dimensions[4].height = 30
    ws1.row_dimensions[5].height = 30

    for (label, value, color), col in zip(kpi_data, kpi_cols):
        ws1.merge_cells(start_row=4, start_column=col, end_row=4, end_column=col+1)
        ws1.merge_cells(start_row=5, start_column=col, end_row=5, end_column=col+1)
        lc = ws1.cell(row=4, column=col, value=label)
        vc = ws1.cell(row=5, column=col, value=value)
        lc.font = Font(bold=True, size=10, color=COLOR["white"])
        lc.fill = PatternFill("solid", fgColor=color)
        lc.alignment = Alignment(horizontal="center", vertical="center")
        vc.font = Font(bold=True, size=14, color=color)
        vc.fill = PatternFill("solid", fgColor=COLOR["gray_bg"])
        vc.alignment = Alignment(horizontal="center", vertical="center")

    # 도메인별 요약 테이블
    ws1.row_dimensions[7].height = 22
    domain_headers = ["감사 도메인", "전체 규칙", "위반 규칙", "위반율(%)", "총 위반 건수"]
    for col, h in enumerate(domain_headers, 1):
        c = ws1.cell(row=7, column=col, value=h)
        header_style(c, bg=COLOR["header_blue"])

    domain_summary = df.groupby("audit_domain").agg(
        전체=("rule_id", "count"),
        위반=("yn_violation", lambda x: (x == "Y").sum()),
        총건수=("violation_count", "sum")
    ).reset_index()
    domain_summary["위반율"] = (domain_summary["위반"] / domain_summary["전체"] * 100).round(1)

    for r_idx, row in enumerate(domain_summary.itertuples(), 8):
        ws1.row_dimensions[r_idx].height = 20
        for c_idx, val in enumerate([row.audit_domain, row.전체, row.위반,
                                      f"{row.위반율}%", row.총건수], 1):
            c = ws1.cell(row=r_idx, column=c_idx, value=val)
            data_style(c, align="center")

    # 총평 (Claude 생성)
    ws1.row_dimensions[12].height = 20
    c = ws1.cell(row=12, column=1, value="■ 경영진 보고용 총평 (AI 생성)")
    c.font = Font(bold=True, size=11, color=COLOR["header_blue"])
    ws1.merge_cells("A12:H12")

    ws1.merge_cells("A13:H18")
    summary_cell = ws1["A13"]
    summary_cell.value = analysis.get("executive_summary", "")
    summary_cell.font = Font(size=10)
    summary_cell.alignment = Alignment(wrap_text=True, vertical="top")
    summary_cell.fill = PatternFill("solid", fgColor="EBF3FB")
    ws1.row_dimensions[13].height = 90

    # 열 너비
    for col, width in [(1,20),(2,10),(3,10),(4,10),(5,15),(6,12),(7,12),(8,15)]:
        set_col_width(ws1, col, width)

    # ────────────────────────────────────────────────────────────
    # 시트 2: 전체 규칙
    # ────────────────────────────────────────────────────────────
    ws2 = wb.create_sheet("전체규칙")
    headers2 = ["규칙ID","규칙명","감사도메인","심각도","출처법령","점검조건","위반건수","위반여부","시정조치"]
    for col, h in enumerate(headers2, 1):
        c = ws2.cell(row=1, column=col, value=h)
        header_style(c)

    for r_idx, row in enumerate(df.itertuples(), 2):
        ws2.row_dimensions[r_idx].height = 18
        vals = [row.rule_id, row.rule_nm, row.audit_domain, row.severity,
                row.source_law, row.condition_desc, row.violation_count,
                row.yn_violation, row.remediation]
        for c_idx, val in enumerate(vals, 1):
            c = ws2.cell(row=r_idx, column=c_idx, value=val)
            data_style(c)
            if c_idx == 4:   # 심각도 컬러
                c.fill = sev_fill.get(str(val), PatternFill())
                c.font = Font(bold=True, color=COLOR["white"], size=10)
                c.alignment = Alignment(horizontal="center", vertical="center")
            if c_idx == 8:   # 위반여부 컬러
                c.fill = viol_fill.get(str(val), PatternFill())
                c.alignment = Alignment(horizontal="center", vertical="center")

    col_widths2 = [8,28,12,10,22,40,10,10,35]
    for col, w in enumerate(col_widths2, 1):
        set_col_width(ws2, col, w)
    ws2.freeze_panes = "A2"

    # ────────────────────────────────────────────────────────────
    # 시트 3: 위반 상세
    # ────────────────────────────────────────────────────────────
    ws3 = wb.create_sheet("위반상세")
    viol_df = df[df["yn_violation"] == "Y"].copy()

    for col, h in enumerate(headers2, 1):
        c = ws3.cell(row=1, column=col, value=h)
        header_style(c, bg="C00000")

    for r_idx, row in enumerate(viol_df.itertuples(), 2):
        ws3.row_dimensions[r_idx].height = 18
        vals = [row.rule_id, row.rule_nm, row.audit_domain, row.severity,
                row.source_law, row.condition_desc, row.violation_count,
                row.yn_violation, row.remediation]
        for c_idx, val in enumerate(vals, 1):
            c = ws3.cell(row=r_idx, column=c_idx, value=val)
            data_style(c)
            if c_idx == 4:
                c.fill = sev_fill.get(str(val), PatternFill())
                c.font = Font(bold=True, color=COLOR["white"], size=10)
                c.alignment = Alignment(horizontal="center", vertical="center")
            if c_idx == 7:
                c.alignment = Alignment(horizontal="center", vertical="center")
                if val and int(val) > 0:
                    c.font = Font(bold=True, color="C00000", size=10)

    for col, w in enumerate(col_widths2, 1):
        set_col_width(ws3, col, w)
    ws3.freeze_panes = "A2"

    # ────────────────────────────────────────────────────────────
    # 시트 4: HIGH 위반
    # ────────────────────────────────────────────────────────────
    ws4 = wb.create_sheet("HIGH위반")
    high_df = df[(df["yn_violation"] == "Y") & (df["severity"] == "HIGH")].copy()

    for col, h in enumerate(headers2, 1):
        c = ws4.cell(row=1, column=col, value=h)
        header_style(c, bg="8B0000")

    for r_idx, row in enumerate(high_df.itertuples(), 2):
        ws4.row_dimensions[r_idx].height = 20
        vals = [row.rule_id, row.rule_nm, row.audit_domain, row.severity,
                row.source_law, row.condition_desc, row.violation_count,
                row.yn_violation, row.remediation]
        for c_idx, val in enumerate(vals, 1):
            c = ws4.cell(row=r_idx, column=c_idx, value=val)
            data_style(c)
            ws4.row_dimensions[r_idx].height = 22
            if c_idx == 4:
                c.fill = sev_fill.get(str(val), PatternFill())
                c.font = Font(bold=True, color=COLOR["white"], size=10)
                c.alignment = Alignment(horizontal="center", vertical="center")

    for col, w in enumerate(col_widths2, 1):
        set_col_width(ws4, col, w)
    ws4.freeze_panes = "A2"

    wb.save(out_path)
    print(f"  Excel 저장 완료 → {out_path}")


# ══════════════════════════════════════════════════════════════
# 3. Word 보고서 생성
# ══════════════════════════════════════════════════════════════
def generate_word(df: pd.DataFrame, analysis: dict, out_path: str):
    """
    구성:
    - 표지
    - 1. 점검 개요
    - 2. 총평 (AI 생성)
    - 3. 도메인별 분석 (AI 생성)
    - 4. 위반 규칙 상세
    - 5. 핵심 시정조치 권고 (AI 생성)
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── 스타일 헬퍼 ──────────────────────────────────────────
    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    def add_heading(text, level=1, color=(31, 78, 121)):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(*color)
        return p

    def add_para(text, bold=False, size=10, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    # ────────────────────────────────────────────────────────────
    # 표지
    # ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IT감사 Rule 엔진\n점검 결과 보고서")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("AI 기반 금융권 IT감사 사전 통제 점검 및 위반 자동 탐지 시스템")
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    doc.add_paragraph()

    # KPI 요약 테이블 (표지용)
    total    = len(df)
    violated = (df["yn_violation"] == "Y").sum()
    high_cnt = ((df["yn_violation"] == "Y") & (df["severity"] == "HIGH")).sum()

    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_labels = ["총 점검 규칙", "위반 탐지", "HIGH 위반", "이상 없음"]
    kpi_values = [f"{total}개", f"{violated}개\n({violated/total*100:.1f}%)",
                  f"{high_cnt}개", f"{total-violated}개"]
    kpi_colors = ["1F4E79", "C00000", "FF4500", "00B050"]

    for col_idx, (label, value, color) in enumerate(zip(kpi_labels, kpi_values, kpi_colors)):
        lc = kpi_table.cell(0, col_idx)
        vc = kpi_table.cell(1, col_idx)
        set_cell_bg(lc, color)
        set_cell_bg(vc, "F2F2F2")
        lc.paragraphs[0].add_run(label).font.color.rgb = RGBColor(255, 255, 255)
        lc.paragraphs[0].runs[0].font.bold = True
        lc.paragraphs[0].runs[0].font.size = Pt(10)
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vc.paragraphs[0].add_run(value).font.bold = True
        vc.paragraphs[0].runs[0].font.size = Pt(13)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"점검 기준일: {datetime.now().strftime('%Y년 %m월 %d일')}").font.size = Pt(11)

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # 1. 점검 개요
    # ────────────────────────────────────────────────────────────
    add_heading("1. 점검 개요", level=1)

    overview_table = doc.add_table(rows=5, cols=2)
    overview_table.style = "Table Grid"
    items = [
        ("점검 기준일",  datetime.now().strftime("%Y년 %m월 %d일")),
        ("점검 대상",    "가상 금융사 (직원 500명 규모, 신용평가사)"),
        ("총 점검 규칙", f"{total}개 (접근통제 42 / 변경관리 15 / 운영통제 13)"),
        ("위반 탐지",    f"{violated}개 ({violated/total*100:.1f}%)"),
        ("점검 근거",    "전자금융감독규정, 금융투자협회 IT감사 가이드라인, 금감원 전자금융감독규정 해설서"),
    ]
    for r_idx, (key, val) in enumerate(items):
        kc = overview_table.cell(r_idx, 0)
        vc = overview_table.cell(r_idx, 1)
        set_cell_bg(kc, "BDD7EE")
        kc.paragraphs[0].add_run(key).font.bold = True
        kc.paragraphs[0].runs[0].font.size = Pt(10)
        vc.paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_paragraph()

    # 도메인별 현황 테이블
    add_heading("도메인별 위반 현황", level=2)
    domain_df = df.groupby("audit_domain").agg(
        전체=("rule_id", "count"),
        위반=("yn_violation", lambda x: (x == "Y").sum()),
        총건수=("violation_count", "sum")
    ).reset_index()
    domain_df["위반율"] = (domain_df["위반"] / domain_df["전체"] * 100).round(1)

    dt = doc.add_table(rows=len(domain_df)+1, cols=5)
    dt.style = "Table Grid"
    d_headers = ["감사 도메인", "전체 규칙", "위반 규칙", "위반율(%)", "총 위반 건수"]
    for c_idx, h in enumerate(d_headers):
        cell = dt.cell(0, c_idx)
        set_cell_bg(cell, "1F4E79")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row in enumerate(domain_df.itertuples(), 1):
        vals = [row.audit_domain, str(row.전체), str(row.위반),
                f"{row.위반율}%", str(row.총건수)]
        for c_idx, val in enumerate(vals):
            cell = dt.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.add_run(val).font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 0:
                set_cell_bg(cell, "F2F2F2")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # 2. 총평 (AI 생성)
    # ────────────────────────────────────────────────────────────
    add_heading("2. 총평", level=1)

    p = doc.add_paragraph()
    p.style.font.size = Pt(10)
    box_run = p.add_run(analysis.get("executive_summary", ""))
    box_run.font.size = Pt(10)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)

    doc.add_paragraph()
    risk_p = doc.add_paragraph()
    risk_run = risk_p.add_run(f"【종합 위험 의견】 {analysis.get('risk_opinion', '')}")
    risk_run.font.size = Pt(10)
    risk_run.font.bold = True
    risk_run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # 3. 도메인별 분석 (AI 생성)
    # ────────────────────────────────────────────────────────────
    add_heading("3. 도메인별 분석", level=1)
    domain_analysis = analysis.get("domain_analysis", {})

    domain_colors = {
        "접근통제": ("C00000", "FFE0E0"),
        "변경관리": ("FF8C00", "FFF3E0"),
        "운영통제": ("1F4E79", "EBF3FB"),
    }

    for domain, text in domain_analysis.items():
        hdr_color, bg_color = domain_colors.get(domain, ("1F4E79", "EBF3FB"))

        # 도메인 헤더
        add_heading(f"3-{list(domain_analysis.keys()).index(domain)+1}. {domain}", level=2)

        # 위반 현황 (간략)
        d_viol = df[(df["audit_domain"] == domain) & (df["yn_violation"] == "Y")]
        d_total = df[df["audit_domain"] == domain]
        stat_p = doc.add_paragraph()
        stat_run = stat_p.add_run(
            f"▶ 점검 {len(d_total)}개 규칙 중 {len(d_viol)}개 위반 탐지 "
            f"({len(d_viol)/len(d_total)*100:.1f}%)"
        )
        stat_run.font.bold = True
        stat_run.font.size = Pt(10)

        # AI 분석 텍스트
        p = doc.add_paragraph()
        p.add_run(text).font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(0.5)

        # 해당 도메인 위반 규칙 테이블
        if len(d_viol) > 0:
            vt = doc.add_table(rows=len(d_viol)+1, cols=4)
            vt.style = "Table Grid"
            v_headers = ["규칙ID", "규칙명", "심각도", "위반 건수"]
            for c_idx, h in enumerate(v_headers):
                cell = vt.cell(0, c_idx)
                set_cell_bg(cell, hdr_color)
                p2 = cell.paragraphs[0]
                run = p2.add_run(h)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            sev_rgb = {"HIGH": (192,0,0), "MEDIUM": (255,140,0), "LOW": (0,176,80)}
            for r_idx, row in enumerate(d_viol.itertuples(), 1):
                vals = [row.rule_id, row.rule_nm, row.severity, str(row.violation_count)]
                for c_idx, val in enumerate(vals):
                    cell = vt.cell(r_idx, c_idx)
                    p2 = cell.paragraphs[0]
                    run = p2.add_run(val)
                    run.font.size = Pt(9)
                    if c_idx == 2:
                        run.font.color.rgb = RGBColor(*sev_rgb.get(str(val), (0,0,0)))
                        run.font.bold = True
                        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if r_idx % 2 == 0:
                        set_cell_bg(cell, "F9F9F9")

        doc.add_paragraph()

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # 4. 핵심 시정조치 권고
    # ────────────────────────────────────────────────────────────
    add_heading("4. 핵심 시정조치 권고", level=1)

    key_actions = analysis.get("key_actions", [])
    action_colors = ["C00000", "FF8C00", "1F4E79", "00B050"]
    for action, color in zip(key_actions, action_colors):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(action)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*tuple(int(color[i:i+2], 16) for i in (0,2,4)))

    doc.add_paragraph()

    # HIGH 위반 시정조치 테이블
    add_heading("HIGH 위반 규칙 시정조치 상세", level=2)
    high_df = df[(df["yn_violation"] == "Y") & (df["severity"] == "HIGH")].copy()

    if len(high_df) > 0:
        ht = doc.add_table(rows=len(high_df)+1, cols=4)
        ht.style = "Table Grid"
        h_headers = ["규칙ID", "규칙명", "위반 건수", "시정조치 방향"]
        for c_idx, h in enumerate(h_headers):
            cell = ht.cell(0, c_idx)
            set_cell_bg(cell, "8B0000")
            p2 = cell.paragraphs[0]
            run = p2.add_run(h)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for r_idx, row in enumerate(high_df.itertuples(), 1):
            vals = [row.rule_id, row.rule_nm, str(row.violation_count), row.remediation]
            for c_idx, val in enumerate(vals):
                cell = ht.cell(r_idx, c_idx)
                p2 = cell.paragraphs[0]
                p2.add_run(val).font.size = Pt(9)
                if r_idx % 2 == 0:
                    set_cell_bg(cell, "FFF0F0")

    doc.save(out_path)
    print(f"  Word  저장 완료 → {out_path}")


# ══════════════════════════════════════════════════════════════
# 메인 실행
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)

    print("=" * 60)
    print("IT감사 보고서 자동 생성 시작")
    print("=" * 60)

    # 1. 데이터 로드
    print("\n[1] violations_summary.csv 로드 중...")
    df = pd.read_csv(SUMMARY_PATH, encoding="utf-8-sig")
    print(f"  총 {len(df)}개 규칙, 위반 {(df['yn_violation']=='Y').sum()}개")

    # 2. Claude API 분석 생성
    print("\n[2] Claude API 분석 생성 중...")
    if _api_key:
        try:
            analysis = generate_analysis(df)
            print("  분석 생성 완료")
        except Exception as e:
            print(f"  API 오류: {e} → 기본 텍스트로 대체")
            analysis = {
                "executive_summary": f"총 {len(df)}개 규칙 점검 결과 {(df['yn_violation']=='Y').sum()}개 위반이 탐지되었습니다. 세부 내용은 하단 도메인별 분석을 참고하시기 바랍니다.",
                "domain_analysis": {
                    "접근통제": "접근통제 분야에서 다수의 위반이 탐지되었습니다.",
                    "변경관리": "변경관리 분야에서 위반이 탐지되었습니다.",
                    "운영통제": "운영통제 분야에서 위반이 탐지되었습니다.",
                },
                "key_actions": ["위반 사항에 대한 즉각적인 시정조치가 필요합니다."],
                "risk_opinion": "전반적인 IT통제 수준에 대한 검토가 필요합니다.",
            }
    else:
        print("  ANTHROPIC_API_KEY 없음 → 기본 텍스트 사용")
        analysis = {
            "executive_summary": f"총 {len(df)}개 규칙 점검 결과 {(df['yn_violation']=='Y').sum()}개 위반이 탐지되었습니다.",
            "domain_analysis": {"접근통제": "-", "변경관리": "-", "운영통제": "-"},
            "key_actions": ["위반 사항 시정조치 필요"],
            "risk_opinion": "상세 분석 필요",
        }

    # 3. 보고서 생성
    date_str = datetime.now().strftime("%Y%m%d")

    print("\n[3] Excel 보고서 생성 중...")
    excel_path = f"{OUT_DIR}/IT감사보고서_{date_str}.xlsx"
    generate_excel(df, analysis, excel_path)

    print("\n[4] Word 보고서 생성 중...")
    word_path = f"{OUT_DIR}/IT감사보고서_{date_str}.docx"
    generate_word(df, analysis, word_path)

    print("\n" + "=" * 60)
    print("보고서 생성 완료!")
    print(f"  Excel → {excel_path}")
    print(f"  Word  → {word_path}")
    print("=" * 60)
